import re
import os
import time
import openai
import pandas
import requests
from .clear_logic import clear_func
from transliterate import translit
from django.conf import settings
from datetime import timedelta
from datetime import datetime
from .models import FileData
from docx import Document
from linteq.secret import OPENAI_TOKEN


def chat_with_gpt(result, langs=None):

    openai.api_key = OPENAI_TOKEN

    if langs:
        dict_doc, _ = create_file_for_gpt(result)
        langs = langs
    else:
        dict_doc, langs = create_file_for_gpt(result)

    res = dict()
    count = 0
    for orig, tran in dict_doc.items():

        prompt = f'Изучив оригинал на языке {langs[0]} сделай пост-редакцию перевода на языке {langs[1]}, исправь ошибки и проверь текст на уникальность. Пришли мне только результат пост-редакции!'
        prompt += f'Оригинал: {orig}; перевод: {tran}'
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": prompt},
            ],
            n=1,
            stop=None,
            temperature=0.0,
            top_p=1.0,
            frequency_penalty=0.0,
        )

        if response and response.choices:
            result = response.choices[0].message.content
            res[orig] = result
        else:
            print('чет не то')
        count += 1

    return res


def create_file_for_gpt(document):

    doc_dict = dict()

    # заголовки языка
    orig_len, tran_len = document.columns.ravel()

    # каждый элемент добавляю в лист и применяю str
    orig_words = list(map(str, document[orig_len].tolist()))
    tran_words = list(map(str, document[tran_len].tolist()))

    # фильтрую лишний row, если нет перевода
    for i in range(len(orig_words)):
        if tran_words[i] != 'nan':
            doc_dict[orig_words[i]] = tran_words[i]

    return doc_dict, [orig_len, tran_len]


def editing_docx(file_path:str, output_folder_path:str, file_name:str):

    # Чтение файла ...
    doc = Document(file_path)
    data = []

    # Перебираем все таблицы в документе
    for table in doc.tables:
        # Перебираем все строки в таблице
        for row in table.rows:
            # Получаем текст каждой ячейки и добавляем в список данных
            row_data = [cell.text for cell in row.cells]
            data.append(row_data)

    orig, tran = data[0][0], data[0][1]

    del data[0]

    # Создаем DataFrame из списка данных
    df = pandas.DataFrame(data)

    # Сохранение файла

    # Создаем новый документ
    doc = Document()

    # Создаем таблицу с двумя колонками
    new_table = doc.add_table(rows=0, cols=2)
    row = new_table.add_row().cells
    row[0].text = orig
    row[1].text = tran

    # Заполняем таблицу значениями из словаря
    for key, value in chat_with_gpt(df, [orig, tran]).items():
        # Добавляем новую строку
        row = new_table.add_row().cells
        # Записываем ключ в первую ячейку
        row[0].text = str(key)
        # Записываем значение во вторую ячейку
        row[1].text = str(value)

    # Сохраняем документ
    doc.save(output_folder_path + file_name)

    return {
        'editing': output_folder_path[6:] + file_name
            }


# Логика рефакторинга документа с расширением xlsx     
def editing_xlsx(file_path:str, output_folder_path:str,
                 file_name:str):

    # Чтение файла ...
    result = pandas.read_excel(file_path)

    # Отправляем файл в chat_with_gpt и сохранение файла
    df = pandas.DataFrame(list(chat_with_gpt(result).items()), columns=result.columns.ravel())
    df.to_excel(output_folder_path + file_name, index=False)

    return {
        'editing': output_folder_path[6:] + file_name
    }


def read_table_file(user_file):

    # Избавление от всех запрещённых символов в названии файла
    illegal_characters = r'[<>:"/\\|?*]'
    filename = re.sub(illegal_characters, '', translit(str(user_file), 'ru', reversed=True))

    # Создание родительской папки
    time_now = re.sub(illegal_characters, '', str(datetime.now()))
    folder_path = f'media/user_request_post_editing/{time_now}'
    
    if os.path.exists(folder_path):
        print(f'[{time_now}] - read_table_file - {folder_path} - Папка уже создана.')
    else:
        os.makedirs(folder_path)
        print(f'[{time_now}] - read_table_file - {folder_path} - Папка была успешно создана.')

    # Создание выходной папки
    output_folder_path = folder_path + '/output/'
    
    if os.path.exists(output_folder_path):
        print(f'[{time_now}] - read_table_file - {output_folder_path} - Выходная папка уже создана.')
    else:
        os.mkdir(output_folder_path)

    # Вненсение пути в бд
    folder_delete_time = datetime.now() + timedelta(
        minutes=settings.STORAGE_TIME)
    
    file_data_model = FileData()
    file_data_model.path = folder_path
    file_data_model.delete_date = folder_delete_time
    file_data_model.save()
    print(f'[{time_now}] - read_table_file - {folder_path} and {folder_delete_time}\
        - Путь и время удаления папки были успешно внесены в базу данных.')
    
    
    # Удаление файлов, у которых истекло время хранения
    clear_func()
    
    
    # Чтение файла
    file = user_file.read()

    # Сохранение файла
    file_path = f'{folder_path}/{filename}'
    
    with open(file_path, 'wb') as f:
        f.write(file)
    print(f'[{time_now}] - read_table_file - {file_path} - Файл был успешно создан.')

    # Получение формата файла и выбор нужной функции
    file_extension = filename.split('.')[-1]
    
    match file_extension:
        
        case 'docx':
            print(f'[{time_now}] - read_table_file - {file_extension}\
                - Начало обработки файла.')
            return editing_docx(file_path=file_path, 
                                output_folder_path=output_folder_path,
                                file_name=filename)

        case 'doc':
            print(f'[{time_now}] - read_table_file - {file_extension}\
                        - Начало обработки файла.')
            return editing_docx(file_path=file_path,
                                output_folder_path=output_folder_path,
                                file_name=filename)
            
        case 'xlsx':
            print(f'[{time_now}] - read_table_file - {file_extension}\
                - Начало обработки файла.')
            return editing_xlsx(file_path=file_path, 
                                output_folder_path=output_folder_path,
                                file_name=filename)
        
        case _:
            return print(f'[{time_now}] - read_table_file - {file_extension}\
                - Этот формат не поддерживается.')
